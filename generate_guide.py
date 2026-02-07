#!/usr/bin/env python3
"""Generate the admin user guide for Piaf restaurant as a .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style definitions ───────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

PIAF_BLUE = RGBColor(0x2a, 0x35, 0x93)
PIAF_YELLOW = RGBColor(0xec, 0xd9, 0x00)
PIAF_GREEN = RGBColor(0x2d, 0x5a, 0x3d)
PIAF_RED = RGBColor(0xdc, 0x35, 0x45)
PIAF_DARK = RGBColor(0x1a, 0x1a, 0x2e)
PIAF_ORANGE = RGBColor(0xe6, 0x7e, 0x22)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Heading styles
for level, size, color in [(1, 22, PIAF_BLUE), (2, 16, PIAF_BLUE), (3, 13, PIAF_DARK)]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    h.paragraph_format.space_after = Pt(8)


def add_colored_paragraph(text, color=None, bold=False, italic=False, size=None, alignment=None):
    """Add a paragraph with specific formatting."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    if color:
        run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_tip_box(text, prefix="Astuce"):
    """Add a tip/info box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"  {prefix} : ")
    run.bold = True
    run.font.color.rgb = PIAF_BLUE
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY
    run2.italic = True
    return p


def add_warning_box(text):
    """Add a warning box."""
    return add_tip_box(text, prefix="Attention")


def add_step(number, text, detail=None):
    """Add a numbered step."""
    p = doc.add_paragraph()
    run = p.add_run(f"  Etape {number} — ")
    run.bold = True
    run.font.color.rgb = PIAF_BLUE
    run2 = p.add_run(text)
    run2.bold = True
    if detail:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1)
        run3 = p2.add_run(detail)
        run3.font.color.rgb = GRAY
        run3.font.size = Pt(10)
    return p


def add_bullet(text, indent=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + indent * 0.8)
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_simple_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2a3593"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="f0f0f8"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


def add_separator():
    """Add a visual separator."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— — —")
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(10)


# ════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PIAF")
run.font.size = Pt(48)
run.font.color.rgb = PIAF_BLUE
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Guide d'utilisation")
run.font.size = Pt(28)
run.font.color.rgb = PIAF_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Interface d'administration des reservations")
run.font.size = Pt(14)
run.font.color.rgb = GRAY

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A destination du personnel du restaurant")
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.italic = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 5.1 — Fevrier 2026")
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("Table des matieres", level=1)

toc_items = [
    ("1.", "Presentation generale"),
    ("2.", "Acceder au tableau de bord"),
    ("3.", "Decouvrir l'interface"),
    ("  3.1", "Vue ordinateur (desktop)"),
    ("  3.2", "Vue mobile"),
    ("4.", "Le calendrier et la navigation par date"),
    ("5.", "Les statistiques du jour"),
    ("6.", "Le plan des tables"),
    ("  6.1", "Comprendre les couleurs des tables"),
    ("  6.2", "Le curseur de creneaux horaires"),
    ("  6.3", "Cliquer sur une table"),
    ("7.", "Gerer les demandes en attente"),
    ("  7.1", "Confirmer une reservation"),
    ("  7.2", "Refuser une reservation"),
    ("  7.3", "Supprimer une reservation"),
    ("8.", "Creer une reservation manuellement"),
    ("9.", "Modifier une reservation existante"),
    ("  9.1", "Modifier le nombre de couverts"),
    ("  9.2", "Modifier la remarque client"),
    ("  9.3", "Modifier l'heure, les tables ou la remarque"),
    ("10.", "Gerer les fermetures exceptionnelles"),
    ("11.", "Les emails automatiques"),
    ("12.", "Configuration du restaurant"),
    ("  12.1", "Horaires d'ouverture et creneaux"),
    ("  12.2", "Capacite des tables"),
    ("13.", "Astuces et bonnes pratiques"),
    ("14.", "En cas de probleme"),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(f"{num}  {title}")
    run.font.size = Pt(10.5)
    if not num.startswith(" "):
        run.bold = True
        run.font.color.rgb = PIAF_BLUE
    else:
        run.font.color.rgb = GRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 1. PRESENTATION GENERALE
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Presentation generale", level=1)

p = doc.add_paragraph()
run = p.add_run("Piaf")
run.bold = True
run.font.color.rgb = PIAF_BLUE
run2 = p.add_run(
    " est le systeme de gestion des reservations du restaurant. "
    "Il permet aux clients de reserver en ligne et a l'equipe de gerer "
    "l'ensemble des reservations depuis un tableau de bord centralise."
)

doc.add_paragraph(
    "Ce guide vous accompagne pas a pas dans l'utilisation de l'interface "
    "d'administration. Vous y apprendrez a :"
)

for item in [
    "Consulter et naviguer dans le calendrier des reservations",
    "Visualiser le plan des tables en temps reel",
    "Confirmer, refuser ou supprimer des demandes de reservation",
    "Creer des reservations manuellement (ex : par telephone)",
    "Modifier les reservations existantes",
    "Gerer les fermetures exceptionnelles du restaurant",
    "Comprendre les emails automatiques envoyes aux clients",
]:
    add_bullet(item)

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 2. ACCEDER AU TABLEAU DE BORD
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Acceder au tableau de bord", level=1)

doc.add_paragraph(
    "Le tableau de bord est accessible depuis un navigateur web "
    "(Chrome, Safari, Firefox...) sur ordinateur, tablette ou telephone."
)

doc.add_heading("Connexion", level=2)

doc.add_paragraph(
    "L'acces est protege par une cle secrete d'administration. "
    "Lorsque vous ouvrez l'interface pour la premiere fois, "
    "le systeme se connecte automatiquement si la cle est configuree."
)

p = doc.add_paragraph()
run = p.add_run("Indicateur de connexion : ")
run.bold = True
p.add_run(
    "en haut de l'ecran, un petit point colore vous indique l'etat de la connexion :"
)

add_simple_table(
    ["Couleur", "Etat", "Signification"],
    [
        ["Vert", "Connecte", "La connexion au serveur est active, tout fonctionne."],
        ["Jaune", "Requete", "Une requete est en cours de traitement."],
        ["Rouge", "Erreur", "La connexion a echoue. Essayez d'actualiser la page."],
    ],
    col_widths=[2.5, 3, 10]
)

add_tip_box(
    "Si le point reste rouge, verifiez votre connexion internet puis "
    "cliquez sur le bouton « Actualiser » en haut a droite de l'ecran."
)

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 3. DECOUVRIR L'INTERFACE
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("3. Decouvrir l'interface", level=1)

doc.add_heading("3.1 Vue ordinateur (desktop)", level=2)

doc.add_paragraph(
    "Sur un ecran large, l'interface est divisee en trois colonnes :"
)

add_simple_table(
    ["Zone", "Position", "Contenu"],
    [
        [
            "Barre laterale gauche",
            "Gauche",
            "Calendrier, statistiques du jour, gestion des fermetures"
        ],
        [
            "Zone centrale",
            "Centre",
            "Plan des tables avec curseur de creneaux horaires, "
            "liste de toutes les reservations du jour, "
            "bouton « Ajouter une reservation »"
        ],
        [
            "Barre laterale droite",
            "Droite",
            "Liste des demandes en attente de confirmation"
        ],
    ],
    col_widths=[4.5, 2.5, 9]
)

p = doc.add_paragraph()
run = p.add_run("En-tete : ")
run.bold = True
p.add_run(
    "tout en haut, vous trouvez le titre « Piaf — Tableau de bord », "
    "l'indicateur de connexion et le bouton « Actualiser »."
)

doc.add_heading("3.2 Vue mobile", level=2)

doc.add_paragraph(
    "Sur telephone ou tablette, l'interface s'adapte avec un affichage "
    "sur une seule colonne. La navigation se fait via une barre en bas "
    "de l'ecran avec quatre onglets :"
)

add_simple_table(
    ["Icone", "Onglet", "Fonction"],
    [
        ["Calendrier", "Calendrier", "Affiche le calendrier et les statistiques"],
        ["Couverts", "Reservations", "Affiche le plan des tables et les reservations du jour"],
        ["Sablier", "En attente", "Affiche les demandes en attente (avec un badge indiquant le nombre)"],
        ["+", "Ajouter", "Ouvre le formulaire de creation de reservation"],
    ],
    col_widths=[2.5, 3, 10.5]
)

add_tip_box(
    "Chaque onglet ouvre un panneau glissant depuis le bas de l'ecran. "
    "Fermez-le en appuyant sur le X ou en touchant la zone grisee en arriere-plan."
)

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 4. CALENDRIER
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("4. Le calendrier et la navigation par date", level=1)

doc.add_paragraph(
    "Le calendrier se trouve dans la barre laterale gauche (ou dans l'onglet "
    "« Calendrier » sur mobile). Il vous permet de selectionner la date "
    "pour laquelle vous souhaitez consulter les reservations."
)

doc.add_heading("Fonctionnement", level=3)

for item in [
    "Utilisez les fleches gauche/droite pour naviguer entre les mois.",
    "Les jours ayant des reservations sont visuellement marques.",
    "Le jour actuel est mis en surbrillance.",
    "Cliquez sur un jour pour charger les reservations de cette date.",
    "Le plan des tables, les statistiques et la liste des reservations "
    "se mettent a jour automatiquement.",
]:
    add_bullet(item)

add_tip_box(
    "Par defaut, le tableau de bord affiche les donnees du jour. "
    "Revenez au jour actuel en cliquant simplement sur la date d'aujourd'hui "
    "dans le calendrier."
)

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 5. STATISTIQUES DU JOUR
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Les statistiques du jour", level=1)

doc.add_paragraph(
    "Juste en dessous du calendrier, la section « Aujourd'hui » affiche "
    "trois indicateurs cles pour la date selectionnee :"
)

add_simple_table(
    ["Indicateur", "Description"],
    [
        ["En attente", "Nombre de reservations qui n'ont pas encore ete confirmees ou refusees."],
        ["Confirmees", "Nombre de reservations confirmees pour cette date."],
        ["Couverts", "Nombre total de personnes attendues (parmi les reservations confirmees)."],
    ],
    col_widths=[4, 12]
)

add_tip_box(
    "Ces chiffres se mettent a jour en temps reel lorsque vous confirmez, "
    "refusez ou ajoutez des reservations."
)

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 6. PLAN DES TABLES
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("6. Le plan des tables", level=1)

doc.add_paragraph(
    "Au centre de l'ecran, le plan des tables offre une vue visuelle "
    "de la salle du restaurant. Il represente les 10 tables du restaurant, "
    "disposees comme dans la salle reelle."
)

doc.add_heading("6.1 Comprendre les couleurs des tables", level=2)

doc.add_paragraph(
    "Chaque table est coloree en fonction de son etat de disponibilite "
    "pour le creneau horaire selectionne :"
)

add_simple_table(
    ["Couleur", "Statut", "Signification"],
    [
        ["Vert", "Libre", "La table est entierement disponible pour ce creneau."],
        ["Orange", "Partielle", "La table est partiellement occupee (certaines places sont encore libres)."],
        ["Rouge", "Complete", "La table est entierement reservee, aucune place disponible."],
        ["Gris (hachures)", "Buffer", "La table est dans une zone tampon (2h avant ou apres une reservation existante)."],
    ],
    col_widths=[3, 2.5, 10.5]
)

p = doc.add_paragraph()
run = p.add_run("Indication de places : ")
run.bold = True
p.add_run(
    "sur chaque table, un indicateur affiche le nombre de places disponibles "
    "par rapport au total, par exemple « 2/4 libres » signifie que 2 places "
    "sur 4 sont encore disponibles."
)

doc.add_heading("6.2 Le curseur de creneaux horaires", level=2)

doc.add_paragraph(
    "Au-dessus du plan des tables se trouve un curseur horizontal (slider). "
    "Il represente les differents creneaux horaires du service de la journee."
)

for item in [
    "Faites glisser le curseur pour voir l'etat des tables a differentes heures.",
    "L'heure actuellement selectionnee est affichee au-dessus du curseur.",
    "Le plan des tables change de couleurs en temps reel selon le creneau choisi.",
]:
    add_bullet(item)

add_tip_box(
    "Ce curseur est particulierement utile pour anticiper les disponibilites "
    "en debut ou en fin de service."
)

doc.add_heading("6.3 Cliquer sur une table", level=2)

doc.add_paragraph(
    "Cliquez sur une table dans le plan pour voir le detail de ses "
    "reservations a ce creneau. Une fenetre s'ouvre avec :"
)

for item in [
    "Le nom du client qui a reserve",
    "L'heure de la reservation",
    "Le nombre de couverts",
    "La possibilite de deplacer la reservation vers une autre table",
    "Un bouton pour modifier la reservation",
]:
    add_bullet(item)

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 7. GERER LES DEMANDES EN ATTENTE
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("7. Gerer les demandes en attente", level=1)

doc.add_paragraph(
    "Lorsqu'un client effectue une reservation en ligne, celle-ci arrive "
    "avec le statut « en attente ». Elle apparait dans la barre laterale "
    "droite (ou l'onglet « En attente » sur mobile)."
)

doc.add_heading("Ce que contient chaque fiche de reservation", level=3)

for item in [
    "Le nom du client (cliquez dessus pour copier l'email)",
    "Le nombre de personnes (cliquez dessus pour le modifier)",
    "La date et l'heure demandees",
    "Le telephone et l'email du client",
    "La remarque du client (allergies, occasion speciale, etc.)",
    "Trois boutons d'action : Confirmer, Refuser, Supprimer",
]:
    add_bullet(item)

# ── 7.1 Confirmer ──
doc.add_heading("7.1 Confirmer une reservation", level=2)

doc.add_paragraph(
    "Confirmer une reservation est l'action principale. Cela attribue "
    "des tables au client et lui envoie un email de confirmation."
)

doc.add_heading("Procedure pas a pas", level=3)

add_step(1, "Cliquez sur le bouton « Confirmer » de la reservation.",
         "La fenetre de confirmation s'ouvre avec le recapitulatif de la demande.")

add_step(2, "Verifiez les informations affichees.",
         "Nom, date, heure, nombre de personnes. "
         "Un indicateur vous montre le nombre de tables necessaires "
         "(ex : « 1 table requise » ou « 2 tables requises »).")

add_step(3, "Selectionnez les tables sur le plan interactif.",
         "Cliquez sur les tables souhaitees : elles se marquent d'une coche. "
         "Les tables selectionnees apparaissent dans un bandeau bleu en bas du plan. "
         "Les couleurs des tables indiquent leur disponibilite (vert = libre, etc.).")

add_step(4, "Cliquez sur « Confirmer avec ces tables ».",
         "La reservation est confirmee, un email est envoye au client, "
         "et les tables sont marquees comme occupees sur le plan.")

add_warning_box(
    "Si un conflit est detecte (table deja reservee a cette heure), "
    "une alerte apparait avec le detail du conflit. Vous pouvez choisir "
    "d'autres tables ou forcer la confirmation."
)

add_tip_box(
    "Vous pouvez ajuster l'heure de la reservation en utilisant le curseur "
    "de creneaux horaires present dans la fenetre de confirmation."
)

# ── 7.2 Refuser ──
doc.add_heading("7.2 Refuser une reservation", level=2)

add_step(1, "Cliquez sur le bouton « Refuser » de la reservation.",
         "La fenetre de refus s'ouvre.")

add_step(2, "Redigez le message a envoyer au client.",
         "Un champ de texte vous permet d'expliquer la raison du refus. "
         "Exemple : « Nous sommes desoles, mais nous sommes complets pour cette date. »")

add_step(3, "Cliquez sur « Envoyer le refus ».",
         "Le client recoit un email avec votre message. "
         "La reservation passe en statut « refusee ».")

# ── 7.3 Supprimer ──
doc.add_heading("7.3 Supprimer une reservation", level=2)

doc.add_paragraph(
    "La suppression est reservee aux cas ou la reservation doit etre "
    "retiree du systeme sans prevenir le client (doublon, test, erreur...)."
)

add_step(1, "Cliquez sur le bouton poubelle de la reservation.")
add_step(2, "Confirmez la suppression dans la fenetre de validation.")

add_warning_box(
    "Cette action est irreversible. La reservation sera definitivement "
    "supprimee et aucun email ne sera envoye au client. "
    "Privilegiez le refus si le client doit etre prevenu."
)

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 8. CREER UNE RESERVATION MANUELLEMENT
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("8. Creer une reservation manuellement", level=1)

doc.add_paragraph(
    "Vous pouvez creer des reservations directement depuis le tableau de "
    "bord, par exemple pour une reservation passee par telephone. "
    "Le bouton « + Ajouter une reservation » se trouve au centre de "
    "l'ecran (ou dans l'onglet « Ajouter » sur mobile)."
)

doc.add_heading("Procedure pas a pas", level=3)

add_step(1, "Cliquez sur « + Ajouter une reservation ».",
         "Le formulaire de creation s'ouvre. Le sous-titre indique : "
         "« Pour les reservations par telephone ».")

add_step(2, "Remplissez les informations obligatoires (marquees d'un *).",
         None)

add_simple_table(
    ["Champ", "Obligatoire", "Description"],
    [
        ["Date", "Oui", "La date de la reservation"],
        ["Nombre de personnes", "Oui", "De 1 a 32 couverts"],
        ["Nom du client", "Oui", "Nom de la personne qui reserve"],
        ["Telephone", "Non", "Numero de telephone du client"],
        ["Email", "Non", "Adresse email du client (pour l'envoi du mail de confirmation)"],
        ["Creneau horaire", "Oui", "Selectionnez l'heure via le curseur"],
        ["Tables", "Oui", "Selectionnez les tables sur le plan interactif"],
        ["Remarque du client", "Non", "Allergies, occasion speciale, chaise bebe, etc."],
        ["Commentaire interne", "Non", "Notes internes pour l'equipe (non visibles par le client)"],
    ],
    col_widths=[4, 2.5, 9.5]
)

add_step(3, "Selectionnez le creneau horaire avec le curseur.")
add_step(4, "Cliquez sur les tables souhaitees dans le plan.",
         "Le systeme vous signale automatiquement tout conflit avec des "
         "reservations existantes.")
add_step(5, "Cliquez sur « Creer la reservation ».",
         "La reservation est creee avec le statut « confirmee ». "
         "Si une adresse email a ete renseignee, un email de confirmation "
         "est envoye au client.")

add_tip_box(
    "Les reservations manuelles recoivent un identifiant commencant par "
    "« MAN- » tandis que les reservations en ligne commencent par « RES- »."
)

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 9. MODIFIER UNE RESERVATION EXISTANTE
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("9. Modifier une reservation existante", level=1)

doc.add_paragraph(
    "Plusieurs modifications sont possibles sur les reservations deja "
    "enregistrees."
)

doc.add_heading("9.1 Modifier le nombre de couverts", level=2)

add_step(1, "Dans la fiche de reservation, cliquez sur le nombre de personnes.",
         "Le nombre est cliquable et affiche une bulle d'aide : "
         "« Cliquer pour modifier le nombre de couverts ».")
add_step(2, "Saisissez le nouveau nombre dans la fenetre qui s'ouvre.",
         "Le champ accepte un nombre entre 1 et 32.")
add_step(3, "Cliquez sur « Valider ».",
         "Le nombre de couverts est mis a jour. Un message de confirmation "
         "s'affiche : « Couverts modifies : X → Y ».")

doc.add_heading("9.2 Modifier la remarque client", level=2)

add_step(1, "Cliquez sur la zone de remarque dans la fiche de reservation.")
add_step(2, "Modifiez le texte dans le champ qui s'ouvre.")
add_step(3, "Cliquez sur « Valider ».",
         "Le message « Remarque mise a jour » confirme la modification.")

doc.add_heading("9.3 Modifier l'heure, les tables ou la remarque", level=2)

doc.add_paragraph(
    "Pour une modification plus complete (changement d'heure ou de table), "
    "utilisez la fenetre d'edition complete :"
)

add_step(1, "Cliquez sur une reservation dans le plan des tables, "
         "puis sur « Modifier la reservation ».",
         "La fenetre d'edition complete s'ouvre.")

add_step(2, "Modifiez les champs souhaites :",
         None)

for item in [
    "Nombre de personnes",
    "Heure (avec les boutons de creneaux disponibles)",
    "Tables (en cliquant sur le plan interactif)",
    "Remarque",
]:
    add_bullet(item)

add_step(3, "Cliquez sur « Enregistrer ».",
         "Les modifications sont sauvegardees et le plan des tables "
         "se met a jour automatiquement.")

add_warning_box(
    "Comme pour la confirmation, le systeme detecte les conflits de tables. "
    "Si un conflit est signale, choisissez une autre table ou un autre creneau."
)

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 10. GERER LES FERMETURES EXCEPTIONNELLES
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("10. Gerer les fermetures exceptionnelles", level=1)

doc.add_paragraph(
    "Vous pouvez bloquer des creneaux horaires pour empecher les clients "
    "de reserver pendant une periode donnee (evenement prive, jour ferie, "
    "travaux, etc.)."
)

doc.add_heading("Ajouter une fermeture", level=3)

add_step(1, "Dans la barre laterale gauche, section « Fermetures », "
         "cliquez sur « + Ajouter une fermeture ».")
add_step(2, "Remplissez le formulaire :",
         None)

for item in [
    "Date de la fermeture",
    "Heure de debut",
    "Heure de fin",
]:
    add_bullet(item)

add_step(3, "Cliquez sur « Ajouter ».",
         "La fermeture apparait dans la liste. Les creneaux concernes "
         "ne seront plus disponibles a la reservation en ligne.")

doc.add_heading("Supprimer une fermeture", level=3)

doc.add_paragraph(
    "Cliquez sur le bouton de suppression a cote de la fermeture concernee "
    "dans la liste. La suppression est immediate."
)

add_tip_box(
    "Les fermetures n'affectent que les nouvelles reservations. "
    "Les reservations deja confirmees pour ces creneaux restent valides."
)

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 11. EMAILS AUTOMATIQUES
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("11. Les emails automatiques", level=1)

doc.add_paragraph(
    "Le systeme envoie plusieurs types d'emails de maniere automatique. "
    "Vous n'avez rien a faire : ils sont declenches par vos actions "
    "dans le tableau de bord."
)

doc.add_heading("Emails envoyes aux clients", level=2)

add_simple_table(
    ["Type d'email", "Quand est-il envoye ?", "Contenu"],
    [
        [
            "Confirmation",
            "Lorsque vous confirmez une reservation",
            "Date, heure, nombre de couverts, coordonnees du restaurant. "
            "Envoye dans la langue du client (FR/EN/NL)."
        ],
        [
            "Refus",
            "Lorsque vous refusez une reservation",
            "Votre message personnalise + coordonnees du restaurant. "
            "Envoye dans la langue du client."
        ],
    ],
    col_widths=[3, 4.5, 8.5]
)

doc.add_heading("Email recapitulatif quotidien", level=2)

doc.add_paragraph(
    "Chaque matin a 8h (sauf lundi et mardi, jours de fermeture), "
    "un email recapitulatif est envoye automatiquement a l'adresse "
    "du restaurant. Cet email contient :"
)

for item in [
    "La liste des reservations confirmees du jour (nom, heure, couverts, tables, contact, remarques)",
    "La liste des reservations en attente",
    "Un tableau clair et organise pour une consultation rapide",
]:
    add_bullet(item)

add_tip_box(
    "Cet email permet de preparer le service en un coup d'oeil, "
    "meme sans ouvrir le tableau de bord."
)

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 12. CONFIGURATION DU RESTAURANT
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("12. Configuration du restaurant", level=1)

doc.add_paragraph(
    "Pour votre information, voici les parametres actuels configures "
    "dans le systeme. Ces informations sont utiles pour comprendre "
    "le fonctionnement des creneaux et des tables."
)

doc.add_heading("12.1 Horaires d'ouverture et creneaux", level=2)

add_simple_table(
    ["Jour", "Service", "Horaires", "Creneaux de reservation"],
    [
        ["Lundi", "Ferme", "—", "—"],
        ["Mardi", "Ferme", "—", "—"],
        ["Mercredi", "Dejeuner", "12h00 - 15h00", "12h00, 12h30, 13h00"],
        ["Jeudi", "Diner", "18h30 - 22h00", "18h30, 19h00, 19h30, 20h00"],
        ["Vendredi", "Diner", "18h30 - 22h30", "18h30, 19h00, 19h30, 20h00, 20h30"],
        ["Samedi", "Brunch", "12h00 - 15h30", "12h00, 12h30, 13h00, 13h30"],
        ["Dimanche", "Brunch", "11h00 - 15h30", "11h00, 11h30, 12h30, 13h00, 13h30"],
    ],
    col_widths=[2.8, 2.5, 3.5, 7.2]
)

doc.add_heading("12.2 Capacite des tables", level=2)

add_simple_table(
    ["Table", "Nombre de places"],
    [
        ["Table 1", "4 places"],
        ["Table 2", "2 places"],
        ["Table 10", "4 places"],
        ["Table 11", "4 places"],
        ["Table 12", "2 places"],
        ["Table 13", "2 places"],
        ["Table 14", "4 places"],
        ["Table 15", "2 places"],
        ["Table 16", "2 places"],
        ["Table 17", "2 places"],
    ],
    col_widths=[4, 4]
)

p = doc.add_paragraph()
run = p.add_run("Total : 10 tables, 28 places assises.")
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Duree d'occupation : ")
run.bold = True
p.add_run(
    "chaque reservation occupe une table pendant 2 heures. "
    "Le systeme bloque automatiquement les 2 heures avant et "
    "les 2 heures apres une reservation sur la meme table "
    "(zone tampon / buffer)."
)

p = doc.add_paragraph()
run = p.add_run("Groupe maximum : ")
run.bold = True
p.add_run("8 personnes (en combinant plusieurs tables).")

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 13. ASTUCES ET BONNES PRATIQUES
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("13. Astuces et bonnes pratiques", level=1)

tips = [
    (
        "Copier l'email d'un client rapidement",
        "Cliquez sur l'icone de copie (presse-papiers) a cote de l'email "
        "dans la fiche de reservation. Le message « Email copie ! » confirme l'action."
    ),
    (
        "Gerer un pic de demandes",
        "Commencez par les reservations les plus anciennes (en haut de la liste). "
        "Utilisez le curseur de creneaux pour verifier rapidement la disponibilite "
        "a differentes heures avant de confirmer."
    ),
    (
        "Utiliser les commentaires internes",
        "Le champ « Commentaire interne » est visible uniquement par l'equipe. "
        "Utilisez-le pour noter des informations de service : "
        "« habitue », « VIP », « allergies graves », etc."
    ),
    (
        "Verifier les conflits",
        "Avant de confirmer, le systeme vous alerte si les tables choisies "
        "sont deja prises. Lisez attentivement le message d'alerte pour "
        "comprendre quel creneau pose probleme."
    ),
    (
        "Preparer le service",
        "Consultez l'email recapitulatif du matin ou ouvrez le tableau de bord "
        "en debut de service. Faites defiler le curseur pour anticiper "
        "les arrivees table par table."
    ),
    (
        "Actualiser les donnees",
        "En cas de doute sur les donnees affichees, cliquez sur "
        "le bouton « Actualiser » en haut a droite pour recharger "
        "toutes les informations depuis le serveur."
    ),
]

for title, desc in tips:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}")
    run.bold = True
    run.font.color.rgb = PIAF_BLUE
    p2 = doc.add_paragraph(desc)
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(10)

add_separator()

# ════════════════════════════════════════════════════════════════════════
# 14. EN CAS DE PROBLEME
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("14. En cas de probleme", level=1)

doc.add_paragraph(
    "Voici les situations les plus courantes et comment y remedier :"
)

problems = [
    (
        "Le point de connexion est rouge",
        "Verifiez votre connexion internet. Cliquez sur « Actualiser ». "
        "Si le probleme persiste, rechargez la page (F5 ou Ctrl+R)."
    ),
    (
        "Je ne vois pas les reservations d'aujourd'hui",
        "Verifiez que la bonne date est selectionnee dans le calendrier. "
        "Si c'est le cas, cliquez sur « Actualiser »."
    ),
    (
        "Un conflit de table est signale",
        "Le systeme vous indique quelle reservation existante entre en conflit. "
        "Choisissez une autre table ou un autre creneau horaire."
    ),
    (
        "J'ai supprime une reservation par erreur",
        "La suppression est irreversible. Si le client avait fourni son email, "
        "contactez-le directement pour proposer une nouvelle reservation."
    ),
    (
        "L'email de confirmation n'a pas ete envoye",
        "L'email n'est envoye que si le client a fourni une adresse email valide. "
        "Verifiez que le champ email est renseigne dans la fiche de reservation."
    ),
    (
        "Le plan des tables ne s'affiche pas correctement",
        "Essayez de selectionner un creneau horaire avec le curseur. "
        "Si le probleme persiste, rechargez la page."
    ),
    (
        "Les fermetures ne bloquent pas les reservations",
        "Les fermetures ne s'appliquent qu'aux nouvelles reservations en ligne. "
        "Les reservations deja confirmees ne sont pas impactees."
    ),
]

for title, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}")
    run.bold = True
    p2 = doc.add_paragraph(desc)
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(10)

add_separator()

# ── Final note ──
for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Bon service !")
run.font.size = Pt(16)
run.font.color.rgb = PIAF_BLUE
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "En cas de question, n'hesitez pas a contacter le responsable technique."
)
run.font.size = Pt(10)
run.font.color.rgb = GRAY
run.italic = True

# ── Save ────────────────────────────────────────────────────────────────

output_path = "/home/user/clients_piaf/Guide_Admin_Piaf.docx"
doc.save(output_path)
print(f"Document genere avec succes : {output_path}")
